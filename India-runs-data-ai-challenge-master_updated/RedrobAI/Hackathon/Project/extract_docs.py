import docx
import os
import sys

base = r'c:\Users\rajan\Downloads\[PUB] India_runs_data_and_ai_challenge\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge'

files = ['README.docx', 'job_description.docx', 'redrob_signals_doc.docx', 'submission_spec.docx']

for fname in files:
    fpath = os.path.join(base, fname)
    outpath = os.path.join(base, fname.replace('.docx', '_extracted.txt'))
    doc = docx.Document(fpath)
    
    with open(outpath, 'w', encoding='utf-8') as out:
        out.write(f"=== {fname} ===\n\n")
        for p in doc.paragraphs:
            out.write(p.text + '\n')
        
        # Also extract tables
        for i, table in enumerate(doc.tables):
            out.write(f"\n--- Table {i+1} ---\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                out.write(' | '.join(cells) + '\n')
    
    print(f"Extracted: {outpath}")
