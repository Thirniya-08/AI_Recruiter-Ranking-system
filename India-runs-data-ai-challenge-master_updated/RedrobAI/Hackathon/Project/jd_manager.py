#!/usr/bin/env python3
"""
jd_manager.py — Job Description Manager for dynamic recruiter-defined JD parsing and skill extraction.
Supports multiple file formats: PDF, DOCX, TXT, JSON, CSV.
"""

import os
import json
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path

try:
    import docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


class JDManager:
    """Manages dynamic Job Descriptions and skill extraction."""
    
    def __init__(self):
        self.current_jd = None
        self.jd_skills = {
            "must_have": [],
            "nice_to_have": [],
            "preferred": []
        }
        self.jd_metadata = {}
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract all text from PDF file."""
        if fitz is None:
            raise ImportError("PyMuPDF (fitz) is not installed. PDF parsing is disabled.")
        
        text = ""
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text() + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF {file_path}: {e}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract all text and tables from DOCX file."""
        if docx is None:
            raise ImportError("python-docx is not installed. DOCX parsing is disabled.")
        
        text = []
        try:
            doc = docx.Document(file_path)
            for p in doc.paragraphs:
                if p.text:
                    text.append(p.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text]
                    if row_cells:
                        text.append(" | ".join(row_cells))
        except Exception as e:
            raise ValueError(f"Error reading DOCX {file_path}: {e}")
        
        return "\n".join(text)
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error reading TXT {file_path}: {e}")
    
    def extract_text_from_json(self, file_path: str) -> str:
        """Extract text from JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, dict):
                return json.dumps(data, indent=2)
            elif isinstance(data, list):
                return "\n".join([json.dumps(item) if isinstance(item, dict) else str(item) for item in data])
            else:
                return str(data)
        except Exception as e:
            raise ValueError(f"Error reading JSON {file_path}: {e}")
    
    def extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file."""
        try:
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                lines = [", ".join(row) for row in reader]
            return "\n".join(lines)
        except Exception as e:
            raise ValueError(f"Error reading CSV {file_path}: {e}")
    
    def parse_jd_file(self, file_path: str) -> Tuple[str, str]:
        """
        Parse JD file and extract text.
        Returns (raw_text, file_type)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            text = self.extract_text_from_pdf(file_path)
            file_type = 'pdf'
        elif file_ext == '.docx':
            text = self.extract_text_from_docx(file_path)
            file_type = 'docx'
        elif file_ext == '.txt':
            text = self.extract_text_from_txt(file_path)
            file_type = 'txt'
        elif file_ext == '.json':
            text = self.extract_text_from_json(file_path)
            file_type = 'json'
        elif file_ext == '.csv':
            text = self.extract_text_from_csv(file_path)
            file_type = 'csv'
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        return text, file_type
    
    def extract_skills_from_jd(self, jd_text: str) -> Dict[str, List[Dict]]:
        """
        Extract skills from JD text using keyword analysis.
        Categorizes into must_have, nice_to_have, and preferred.
        """
        jd_text_lower = jd_text.lower()
        
        # Skill categories
        must_have_keywords = {
            'python': ['python', 'python3', 'python 3', 'py'],
            'machine learning': ['machine learning', 'ml', 'deep learning', 'neural', 'tensorflow', 'pytorch', 'keras'],
            'data analysis': ['data analysis', 'analytics', 'data science', 'data scientist'],
            'sql': ['sql', 'mysql', 'postgresql', 'database'],
            'distributed systems': ['distributed', 'spark', 'hadoop', 'kafka'],
            'cloud': ['aws', 'azure', 'gcp', 'cloud', 'ec2', 's3', 'lambda'],
            'docker': ['docker', 'kubernetes', 'k8s', 'container'],
            'api development': ['api', 'rest', 'fastapi', 'flask', 'django', 'endpoint'],
        }
        
        nice_to_have_keywords = {
            'nlp': ['nlp', 'nlp', 'natural language', 'transformer', 'bert', 'gpt'],
            'computer vision': ['computer vision', 'cv', 'image', 'object detection', 'opencv'],
            'reinforcement learning': ['reinforcement learning', 'rl', 'dqn', 'policy gradient'],
            'data engineering': ['data engineering', 'etl', 'pipeline', 'data warehouse'],
            'big data': ['big data', 'petabyte', 'terabyte'],
        }
        
        preferred_keywords = {
            'agile': ['agile', 'scrum', 'kanban'],
            'git': ['git', 'github', 'gitlab', 'bitbucket'],
            'ci/cd': ['ci/cd', 'jenkins', 'gitlab-ci', 'github actions', 'devops'],
            'monitoring': ['monitoring', 'prometheus', 'grafana', 'logging', 'splunk'],
        }
        
        skills = {
            'must_have': [],
            'nice_to_have': [],
            'preferred': []
        }
        
        # Extract must_have skills
        for skill, keywords in must_have_keywords.items():
            found_keywords = [kw for kw in keywords if kw in jd_text_lower]
            if found_keywords:
                skills['must_have'].append({
                    'label': skill,
                    'keywords': found_keywords,
                    'description': f"Mentioned in JD: {', '.join(set(found_keywords))}"
                })
        
        # Extract nice_to_have skills
        for skill, keywords in nice_to_have_keywords.items():
            found_keywords = [kw for kw in keywords if kw in jd_text_lower]
            if found_keywords:
                skills['nice_to_have'].append({
                    'label': skill,
                    'keywords': found_keywords,
                    'description': f"Mentioned in JD: {', '.join(set(found_keywords))}"
                })
        
        # Extract preferred skills
        for skill, keywords in preferred_keywords.items():
            found_keywords = [kw for kw in keywords if kw in jd_text_lower]
            if found_keywords:
                skills['preferred'].append({
                    'label': skill,
                    'keywords': found_keywords,
                    'description': f"Mentioned in JD: {', '.join(set(found_keywords))}"
                })
        
        return skills
    
    def extract_metadata_from_jd(self, jd_text: str) -> Dict:
        """Extract metadata like job title, company, requirements from JD."""
        metadata = {
            'job_title': None,
            'company': None,
            'experience_required': None,
            'location': None,
            'employment_type': None,
            'salary_range': None,
        }
        
        jd_lower = jd_text.lower()
        
        # Try to extract job title (usually in first line or with "Position:" prefix)
        title_patterns = [
            r'(?:position|job title|role):\s*([^\n]+)',
            r'^([^\n]{10,100})\n',
        ]
        for pattern in title_patterns:
            match = re.search(pattern, jd_text, re.IGNORECASE)
            if match:
                metadata['job_title'] = match.group(1).strip()
                break
        
        # Extract experience requirement
        exp_pattern = r'(\d+)\+?\s*(?:years|yrs|years of experience|years experience|yoe)'
        match = re.search(exp_pattern, jd_lower)
        if match:
            metadata['experience_required'] = int(match.group(1))
        
        # Extract employment type
        for emp_type in ['full-time', 'part-time', 'contract', 'freelance', 'temporary']:
            if emp_type in jd_lower:
                metadata['employment_type'] = emp_type
                break
        
        return metadata
    
    def load_jd(self, jd_text: str, file_type: str = 'raw') -> Dict:
        """Load and parse a Job Description."""
        self.current_jd = jd_text
        
        # Extract skills
        self.jd_skills = self.extract_skills_from_jd(jd_text)
        
        # Extract metadata
        self.jd_metadata = self.extract_metadata_from_jd(jd_text)
        
        return {
            'status': 'loaded',
            'file_type': file_type,
            'jd_text_length': len(jd_text),
            'skills': self.jd_skills,
            'metadata': self.jd_metadata
        }
    
    def get_jd_summary(self) -> Dict:
        """Get current JD summary."""
        if not self.current_jd:
            return {'error': 'No JD loaded'}
        
        return {
            'job_title': self.jd_metadata.get('job_title', 'Not specified'),
            'company': self.jd_metadata.get('company', 'Not specified'),
            'experience_required': self.jd_metadata.get('experience_required', 'Not specified'),
            'must_have_skills': len(self.jd_skills['must_have']),
            'nice_to_have_skills': len(self.jd_skills['nice_to_have']),
            'preferred_skills': len(self.jd_skills['preferred']),
            'skills': self.jd_skills,
            'metadata': self.jd_metadata
        }
    
    def clear_jd(self):
        """Clear current JD."""
        self.current_jd = None
        self.jd_skills = {
            'must_have': [],
            'nice_to_have': [],
            'preferred': []
        }
        self.jd_metadata = {}


# Global JD manager instance
jd_manager = JDManager()
